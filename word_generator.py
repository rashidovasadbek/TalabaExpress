from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random
import os

PAGE_TITLES = {
    'uz': {
        'reja_sarlavha': "REJA",
        'kirish': "1. Kirish",
        'xulosa': "4. Xulosa",
        'adabiyotlar': "5. Foydalanilgan adabiyotlar ro'yxati",
        'adabiyotlar_title': "FOYDALANILGAN ADABIYOTLAR RO'YXATI",
        'label_study_year': "o'quv yili",
        
         
        # TITUL sahifasi uchun yorliqlar (Labels)
        'label_topic': "MAVZU", 
        'label_bajaruvchi': "Bajardi",
        'label_group':"Guruh",
        'label_rahbar': "Ilmiy rahbar",
        'label_study_year': "o'quv yili",
        'title_referat': "REFERAT",
        'title_mustaqil_ish': "MUSTAQIL ISH",
        'vazirlik_title': "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI",# Yangi qo'shilgan kalit
    },
    'ru': {
        'reja_sarlavha': "ПЛАН",
        'kirish': "1. ВВЕДЕНИЕ",
        'xulosa': "4. ЗАКЛЮЧЕНИЕ",
        'adabiyotlar': "5. Список использованной литературы",
        'adabiyotlar_title': "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ",
        
         # ТИТУЛЬНЫЙ ЛИСТ uchun yorliqlar (Labels)
        'label_topic': "ТЕМА",
        'label_bajaruvchi': "Выполнил",
        'label_group':"Группа",
        'label_rahbar': "Научный руководитель",
        'label_study_year': "учебный год",
        'title_referat': "РЕФЕРАТ",
        'title_mustaqil_ish': "САМОСТОЯТЕЛЬНАЯ РАБОТА",
        'vazirlik_title': "МИНИСТЕРСТВО ВЫСШЕГО ОБРАЗОВАНИЯ, НАУКИ И ИННОВАЦИЙ РЕСПУБЛИКИ УЗБЕКИСТАН",
    },
    'en': {
        'reja_sarlavha': "CONTENTS",
        'kirish': "1. INTRODUCTION",
        'xulosa': "4. CONCLUSION",
        'adabiyotlar': "5. List of References",
        'adabiyotlar_title': "LIST OF REFERENCES",
        
         # TITLE PAGE uchun yorliqlar (Labels)
        'label_topic': "TOPIC",
        'label_bajaruvchi': "Completed by",
        'label_group':"Group",
        'label_rahbar': "Supervisor",
        'label_study_year': "academic year",
        'title_referat': "REPORT",
        'title_mustaqil_ish': "TERM PAPER",
        'vazirlik_title': "MINISTRY OF HIGHER EDUCATION, SCIENCE AND INNOVATION OF THE REPUBLIC OF UZBEKISTAN",# Yangi qo'shilgan kalit
    },
    'kr': {
        'reja_sarlavha': "РЕЖА",
        'kirish': "1. Кириш",
        'xulosa': "4. Хулоса",
        'adabiyotlar': "5. Фойдаланилган адабиётлар рўйхати",
        'adabiyotlar_title': "ФОЙДАЛАНИЛГАН АДАБИЁТЛАР РЎЙХАТИ",
        
          # ТИТУЛ sahifasi uchun yorliqlar (Labels)
        'label_topic': "МАВЗУ",
        'label_bajaruvchi': "Бажарди",
        'label_group':"Группа",
        'label_rahbar': "Илмий раҳбар",
        'label_study_year': "ўқув йили",
        'title_referat': "РЕФЕРАТ",
        'title_mustaqil_ish': "МУСТАҚИЛ ИШ",
        'vazirlik_title': "ЎЗБЕКИСТОН РЕСПУБЛИКАСИ ОЛИЙ ТАЪЛИМ, ФАН ВА ИННОВАЦИЯЛАР ВАЗИРЛИГИ",
    },
}

def set_font_style(run, font_name="Times New Roman", size=14, is_bold=False, is_italic=False):
    """Run obyekti uchun shrift stilini o'rnatadi."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = is_bold
    run.italic = is_italic

def set_paragraph_alignment(p, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Paragrafning tekislash usulini (Chap/O'rta/Keng) o'rnatadi."""
    p.alignment = alignment

def create_title_page(document: Document, doc_data: dict, title_data: dict):
    """
    Standartga muvofiq Titul sahifasini yaratadi.
    doc_data: Talabaning kiritilgan ma'lumotlari (FIO, Mavzu kabi)
    title_data: Tilga moslashtirilgan rasmiy matnlar
    """
    
    # --- Talabaning shaxsiy ma'lumotlari ---
    topic = doc_data.get('topic', 'Mavzu nomi kiritilmagan')
    uni_kafedra = doc_data.get('uni_kafedra', 'Fakultet va Kafedra nomi')
    student_fio = doc_data.get('student_fio', 'Talaba F.I.O.si kiritilmagan') 
    student_group = doc_data.get('student_group', 'Guruh raqami kiritilmagan')
 
    p_vazir = document.add_paragraph(title_data['vazirlik_title']) # <<< YANGI KALIT
    set_paragraph_alignment(p_vazir, WD_ALIGN_PARAGRAPH.CENTER)
    set_font_style(p_vazir.runs[0], size=20, is_bold=False)
    document.add_paragraph() 
    
    p_uni = document.add_paragraph(uni_kafedra.upper())
    set_paragraph_alignment(p_uni, WD_ALIGN_PARAGRAPH.CENTER)
    set_font_style(p_uni.runs[0], size=20, is_bold=False)
    
    work_title_text = title_data.get('work_title_display', 'REFERAT').upper()   
    p_work_type = document.add_paragraph()
    run_work_type = p_work_type.add_run(work_title_text) 
    set_paragraph_alignment(p_work_type, WD_ALIGN_PARAGRAPH.CENTER)
    
    raw_work_type = doc_data.get('raw_work_type', 'referat')
    if raw_work_type == 'mustaqil_ish':
        set_font_style(run_work_type, size=36, is_bold=True)
    else:
        set_font_style(run_work_type, size=60, is_bold=True)


    for _ in range(2): 
        document.add_paragraph()
        
    p_topic = document.add_paragraph()
    set_paragraph_alignment(p_topic, WD_ALIGN_PARAGRAPH.CENTER)
    
    p_topic.add_run(title_data['label_topic'] + ": ")
    set_font_style(p_topic.runs[0], size=16, is_bold=True) 
    
    p_topic.add_run(topic.upper()) 
    set_font_style(p_topic.runs[1] , size=16, is_bold=True) 
    
    for _ in range(2):
        document.add_paragraph()
        
    p_bajardi = document.add_paragraph()
    set_paragraph_alignment(p_bajardi, WD_ALIGN_PARAGRAPH.LEFT)
    p_bajardi.add_run(title_data['label_bajaruvchi'] + ": ")
    set_font_style(p_bajardi.runs[0], size=14, is_bold=False)
    p_bajardi.add_run(student_fio)
    set_font_style(p_bajardi.runs[1], size=14, is_bold=True) 
    
    p_group = document.add_paragraph()
    set_paragraph_alignment(p_group, WD_ALIGN_PARAGRAPH.LEFT)
    p_group.add_run(title_data['label_group'] +": ")
    set_font_style(p_group.runs[0], size=14, is_bold=False)
    p_group.add_run(student_group)
    set_font_style(p_group.runs[1], size=14, is_bold=True)
    
    
    p_rahbar = document.add_paragraph()
    set_paragraph_alignment(p_rahbar, WD_ALIGN_PARAGRAPH.LEFT)
    
    p_rahbar.add_run(title_data['label_rahbar'] + ": ")
    set_font_style(p_rahbar.runs[0], size=14, is_bold=False)
    
    p_rahbar.add_run("") 
    set_font_style(p_rahbar.runs[1], size=14, is_bold=True)
    
    for _ in range(3): 
        document.add_paragraph()

    p_year = document.add_paragraph(title_data['city_year'])
    set_paragraph_alignment(p_year, WD_ALIGN_PARAGRAPH.CENTER)
    set_font_style(p_year.runs[0], size=12, is_bold=True)
    p_year.add_run(" " + title_data['label_study_year']) 
    set_font_style(p_year.runs[1], size=12, is_bold=False)  
    
def add_referat_reja(document, ai_reja_data, set_font_style, set_paragraph_alignment, title_data: dict):
    
    document.add_page_break()
    
    p_reja_title = document.add_paragraph(title_data['reja_sarlavha'])
    set_paragraph_alignment(p_reja_title, WD_ALIGN_PARAGRAPH.CENTER)
    set_font_style(p_reja_title.runs[0], size=18, is_bold=True)
    document.add_paragraph() 
        
    reja_qismlari = []
    
    reja_qismlari.append((title_data['kirish'], "Kirish"))
    
    main_sections_count = len(ai_reja_data) if ai_reja_data and isinstance(ai_reja_data, list) else 0

    if main_sections_count > 0:
        
        for i, bob_data in enumerate(ai_reja_data): 
            main_title = bob_data.get("main_title", f"Asosiy Bob {i+2}") 
            sub_titles = bob_data.get("sub_titles", [])
            
            bob_raqami = i + 2 
            
            reja_qismlari.append((f"{bob_raqami}. {main_title}", f"Bob_{i+1}"))
            
            for j, sub_title in enumerate(sub_titles): 
                sub_band_raqami = j + 1
                reja_qismlari.append((f"{bob_raqami}.{sub_band_raqami}. {sub_title}", f"Bob_{i+1}_band_{sub_band_raqami}"))
       
    
        xulosa_raqami = main_sections_count + 2
        adabiyotlar_raqami = main_sections_count + 3  
        
        xulosa_sarlavhasi = title_data['xulosa'].split('. ', 1)[-1]      
        
        reja_qismlari.append((f"{xulosa_raqami}. {xulosa_sarlavhasi}", "Xulosa"))
        
        adabiyotlar_sarlavhasi = title_data['adabiyotlar'].split('. ', 1)[-1]

        reja_qismlari.append((f"{adabiyotlar_raqami}. {adabiyotlar_sarlavhasi}", "Adabiyotlar"))
    else:
        reja_qismlari.append((title_data['xulosa'], "Xulosa")) 
        reja_qismlari.append((title_data['adabiyotlar'], "Adabiyotlar")) 
        
        
    for text, slug in reja_qismlari:
        p_reja_band = document.add_paragraph(text)
        set_paragraph_alignment(p_reja_band, WD_ALIGN_PARAGRAPH.LEFT)
        
        if text.count('.') <= 1:
            set_font_style(p_reja_band.runs[0], size=14, is_bold=True)
        else:
            set_font_style(p_reja_band.runs[0], size=14, is_bold=False)
            
    document.add_page_break()

def add_content_section(document, title: str, content: str, set_font_style, set_paragraph_alignment):
    """
    Hujjatga formatlangan bob sarlavhasi va matnini qo'shadi.
    """
    
    p_title = document.add_paragraph(title)
    set_paragraph_alignment(p_title, WD_ALIGN_PARAGRAPH.CENTER)
    set_font_style(p_title.runs[0], size=14, is_bold=True)
    
    document.add_paragraph() 
    
    for line in content.split('\n'):
        if line.strip():
            p_content = document.add_paragraph(line)
            set_paragraph_alignment(p_content, WD_ALIGN_PARAGRAPH.JUSTIFY)
            set_font_style(p_content.runs[0], size=14, is_bold=False)
   
def generate_word_file(work_type:str, ai_reja_data, intro_text:str, conclusion_text:str, references_text:str, doc_data: dict, title_data: dict) -> str:
    
    document = Document()
    
    create_title_page(document, doc_data, title_data)
    
    work_type_lower = work_type.lower().replace('_', ' ')
    if work_type_lower.lower() in ['refarat', 'mustaqil ish']:
        
        add_referat_reja(
             document=document, 
             ai_reja_data=ai_reja_data, 
             title_data=title_data,
             set_font_style=set_font_style, 
             set_paragraph_alignment=set_paragraph_alignment
        )
        
        add_content_section(
            document=document,
            title=title_data['title_intro'], 
            content=intro_text, 
            set_font_style=set_font_style,
            set_paragraph_alignment=set_paragraph_alignment
        )
        document.add_page_break()
    
    for i, bob_data in enumerate(ai_reja_data):
        bob_raqami = i + 2 
        main_title = bob_data["main_title"]
        content = bob_data["content"] 
        
        add_content_section(
            document=document,
            title=f"{bob_raqami}. {main_title}",
            content=content,
            set_font_style=set_font_style,
            set_paragraph_alignment=set_paragraph_alignment
        )
        document.add_page_break() 
        
    add_content_section(
        document=document,
        title=title_data['title_conclusion'], 
        content=conclusion_text,
        set_font_style=set_font_style,
        set_paragraph_alignment=set_paragraph_alignment
    )
    document.add_page_break() 
    
    add_content_section(
        document=document,
        title=title_data['title_references'],
        content=references_text,
        set_font_style=set_font_style,
        set_paragraph_alignment=set_paragraph_alignment
    )
    
    student_info_raw = doc_data.get('student_fio')
    
    fio_parts = student_info_raw.strip().split()
    
    fio_base = 'Nomalum_Talaba'
    if len(fio_parts) >= 2:
        fio_base = f"{fio_parts[0]}_{fio_parts[1]}" 
    elif len(fio_parts) == 1:
        fio_base = fio_parts[0]
        
    fio_base = fio_base.replace('-', '_')
    
    work_type_name = work_type.replace(' ', '_').capitalize()
    
    file_name = f"{fio_base}_{work_type_name}.docx"
    
    file_path = os.path.join(os.getcwd(), file_name)
    document.save(file_path)
    
    return file_path